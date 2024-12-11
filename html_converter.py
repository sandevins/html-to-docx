import sys
import base64
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image

def convert_html_to_docx(input_file, output_file):
    try:
        # Cargar el archivo HTML
        with open(input_file, "r", encoding="utf-8") as file:
            soup = BeautifulSoup(file, "html.parser")

        # Crear un documento DOCX
        doc = Document()

        # Procesar el contenido HTML
        for element in soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6", "table", "img"]):
            if element.name.startswith("h"):  # Encabezados
                doc.add_heading(element.get_text(), level=int(element.name[1]))
            elif element.name == "p":  # Párrafos
                doc.add_paragraph(element.get_text())
            elif element.name == "table":  # Tablas
                table = doc.add_table(rows=1, cols=len(element.find("tr").find_all("td") or element.find("tr").find_all("th")))
                table.style = 'Table Grid'
                # Procesar filas de la tabla
                for row in element.find_all("tr"):
                    cells = row.find_all(["td", "th"])
                    doc_row = table.add_row().cells if table.rows else table.rows[0].cells
                    for i, cell in enumerate(cells):
                        doc_row[i].text = cell.get_text(strip=True)
            elif element.name == "img":  # Imágenes
                img_src = element.get("src", "")
                if img_src.startswith("data:image/"):
                    # Decodificar imagen en Base64
                    header, encoded = img_src.split(",", 1)
                    img_data = base64.b64decode(encoded)
                    image_stream = BytesIO(img_data)
                    img = Image.open(image_stream)
                    img_format = img.format.lower()
                    if img_format in ["jpeg", "png", "bmp", "gif"]:
                        image_stream.seek(0)
                        doc.add_picture(image_stream, width=Inches(2))  # Ajustar el tamaño según sea necesario

        # Guardar el archivo como DOCX
        doc.save(output_file)
        print(f"Archivo convertido con éxito: {output_file}")
    except Exception as e:
        print(f"Error al convertir el archivo: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Uso: python script.py <archivo_entrada.html> <archivo_salida.docx>")
    else:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
        convert_html_to_docx(input_file, output_file)
